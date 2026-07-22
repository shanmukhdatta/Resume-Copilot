"""Renders the final resume into a DOCX file using python-docx."""
import os

from app.config import get_settings
from app.schemas.resume_schema import GeneratedResume
from app.utils.file_utils import ensure_dir


def render_docx(resume: GeneratedResume, session_id: str) -> str:
    import docx
    from docx.shared import Pt

    settings = get_settings()
    document = docx.Document()

    title = document.add_heading(resume.contact.full_name or "Resume", level=0)

    contact_bits = [b for b in (resume.contact.email, resume.contact.phone, resume.contact.location) if b]
    if contact_bits:
        document.add_paragraph(" | ".join(contact_bits))

    if resume.summary:
        document.add_heading("Summary", level=1)
        document.add_paragraph(resume.summary)

    if resume.skills:
        document.add_heading("Skills", level=1)
        for cat in resume.skills:
            skills = cat.get("skills", []) if isinstance(cat, dict) else cat.skills
            category = cat.get("category", "") if isinstance(cat, dict) else cat.category
            document.add_paragraph(f"{category}: {', '.join(skills)}")

    if resume.experience:
        document.add_heading("Experience", level=1)
        for exp in resume.experience:
            e = exp if isinstance(exp, dict) else exp.model_dump()
            p = document.add_paragraph()
            run = p.add_run(f"{e.get('title', '')}, {e.get('company', '')}")
            run.bold = True
            document.add_paragraph(f"{e.get('start_date', '')} - {e.get('end_date', '')}")
            for bullet in e.get("bullets", []):
                document.add_paragraph(bullet, style="List Bullet")

    if resume.projects:
        document.add_heading("Projects", level=1)
        for proj in resume.projects:
            p = proj if isinstance(proj, dict) else proj.model_dump()
            run = document.add_paragraph().add_run(p.get("name", ""))
            run.bold = True
            for bullet in p.get("bullets", []):
                document.add_paragraph(bullet, style="List Bullet")

    if resume.education:
        document.add_heading("Education", level=1)
        for edu in resume.education:
            e = edu if isinstance(edu, dict) else edu.model_dump()
            document.add_paragraph(f"{e.get('degree', '')}, {e.get('institution', '')} ({e.get('start_date', '')} - {e.get('end_date', '')})")

    out_dir = ensure_dir(settings.EXPORT_DIR)
    out_path = os.path.join(out_dir, f"{session_id}_resume.docx")
    document.save(out_path)
    return out_path
